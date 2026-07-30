from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def heading1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs: r.font.color.rgb = RGBColor(0,0,0)
def heading2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.color.rgb = RGBColor(0,0,0)
def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(12); r.font.name = 'Times New Roman'; r.bold = bold
    return p
def lp(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Paragraph'); p.clear()
    if bold_prefix:
        rb = p.add_run(bold_prefix + ' '); rb.bold = True; rb.font.size = Pt(12); rb.font.name = 'Times New Roman'
    r = p.add_run(text); r.font.size = Pt(12); r.font.name = 'Times New Roman'

heading1('INVENTION DISCLOSURE FORM (IDS)')
doc.add_paragraph()
para('Date: ____/ 2026')
para('Client: IIT BHILAI [IITBH]')
doc.add_paragraph()

heading2('Please provide an appropriate Title of the Invention:')
para('System and Method for Dynamic Spatial Repositioning, Orientation Adjustment, and Transmission Parameter Adaptation of Movable Wireless Communication Nodes Using Optimization-Driven Control Under Changing Environmental Conditions')
doc.add_paragraph()

para('Technology Domain of the Invention (Please choose relevant)', bold=True)
para('Wireless Telecommunications (all generations: 4G, 5G, 5G-Advanced, 6G, and any future standard), Radio-Frequency / Microwave / Millimeter-Wave / Terahertz / Optical Wireless / Visible Light Communication, Swarm-Intelligence and Population-Based Optimization, Machine Learning and Artificial Intelligence for Network Management, Cognitive and Adaptive Networking, Robotic and Autonomous Mobility Systems, Dynamic Wireless Infrastructure Management, Internet of Things (IoT), Smart Environments.')
doc.add_paragraph()

heading2('Background of the Invention')
para('Please clearly discuss the problems the invention tries to solve. Please list as many.', bold=True)
doc.add_paragraph()
para('Wireless communication networks across all frequency bands\u2014from sub-1 GHz through microwave, millimeter-wave (mmWave, 24\u2013100 GHz), terahertz (THz, 100 GHz\u201310 THz), infrared, and optical frequencies\u2014face fundamental physical and environmental challenges that degrade service quality:')

probs = [
    ('1.', 'Propagation losses that increase with carrier frequency, including free-space path loss (FSPL), atmospheric absorption, rain attenuation, foliage loss, and diffraction losses.'),
    ('2.', 'Susceptibility to static physical blockages\u2014walls, partitions, furniture, structural elements, terrain features, buildings, vehicles, and any other fixed or semi-fixed obstacles composed of materials with varying radio-frequency attenuation characteristics.'),
    ('3.', 'Susceptibility to dynamic blockages\u2014humans, animals, moving vehicles, robotic platforms, drones, equipment, machinery, environmental phenomena (rain, fog, dust, smoke, vegetation movement), and any other time-varying obstruction that degrades, reflects, scatters, or absorbs the signal.'),
    ('4.', 'Time-varying user density and traffic demand distributions, where users congregate in different spatial zones at different times, creating shifting coverage requirements that static infrastructure cannot serve optimally.'),
    ('5.', 'Narrow beamwidths (especially at higher frequencies) that limit spatial coverage per node, requiring precise beam alignment that is disrupted by any change in environment or user position.'),
    ('6.', 'Inter-node interference when multiple transmitting nodes serve overlapping areas, degrading SINR and throughput.'),
    ('7.', 'Fixed transmitting node installations cannot reposition to: (a) adapt to new user density zones; (b) clear new blockages; (c) exploit vertical or angular geometry changes; (d) reduce inter-node interference through spatial separation; or (e) minimize path loss by reducing distance to users.'),
    ('8.', 'Electronic-only beam adaptation (beamforming, beam steering, beam switching) cannot overcome scenarios where no viable propagation path (direct, reflected, or diffracted) exists from the fixed node position to the user.'),
    ('9.', 'Existing mobile relay/drone-based approaches lack a structured decision framework for determining when physical movement is justified versus when electronic adaptation suffices, leading to unnecessary mechanical actuation, latency, energy waste, and wear.'),
    ('10.', 'No existing solution provides a unified, hierarchical, optimization-driven framework that jointly considers: (a) electronic parameter adaptation; (b) physical repositioning in one or more spatial dimensions; (c) time-varying environmental sensing; (d) multi-node coordination; and (e) operator-configurable quality-of-service objectives\u2014in a single integrated system.'),
]
for bp, txt in probs:
    lp(txt, bold_prefix=bp)
doc.add_paragraph()

para('Please list already existing prior art / patents / methods / products (if you are aware of any)', bold=True)
doc.add_paragraph()
table = doc.add_table(rows=6, cols=3); table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Prior art No.', 'Publication Date', 'Brief Description of the Prior art']):
    c = table.rows[0].cells[i]; c.text = h
    for p in c.paragraphs:
        for r in p.runs: r.bold = True; r.font.size = Pt(11)
rows = [
    ['US20230292199A1', 'Sept 14, 2023', 'Fixed small cell deployments\u2014high density, high cost, no adaptation to dynamic blockages.'],
    ['WO2024096638A1', 'May 10, 2024', 'Electronic beamforming\u2014cannot overcome complete physical blockages from fixed position.'],
    ['WO2018071453A1', 'April 19, 2018', 'UAV relays\u2014outdoor macro coverage, lacks indoor precision, airspace constraints.'],
    ['RIS (various)', 'Various', 'Passive reflection only\u2014cannot relocate source, cannot generate independent coverage.'],
    ['Mobile relays (various)', 'Various', 'No hierarchical decision framework for move-vs-adapt, leading to excess actuation.'],
]
for ri, rd in enumerate(rows):
    for ci, cd in enumerate(rd): table.rows[ri+1].cells[ci].text = cd
doc.add_paragraph()

heading2('Invention description')
para('Briefly summarize how the aforementioned problem is solved by the invention?', bold=True)
doc.add_paragraph()
para('The invention provides a system and method for dynamically repositioning, reorienting, and adapting the transmission parameters of one or more movable wireless communication nodes\u2014where "wireless communication node" encompasses any device or apparatus\u2014existing, under development, or to be developed in the future\u2014capable of transmitting, receiving, relaying, repeating, amplifying, reflecting, or otherwise processing wireless signals, including but not limited to: access points, base stations (macro, micro, pico, femto), routers, relay nodes, repeaters, range extenders, distributed antenna system (DAS) heads, radio units (RUs), remote radio heads (RRHs), reconfigurable intelligent surfaces with active elements, intelligent reflecting surfaces, backhaul/fronthaul nodes, mesh network nodes, ad-hoc network nodes, mobile hotspots, Wi-Fi extenders, IoT gateways, LoRa gateways, satellite ground terminals, vehicle-to-everything (V2X) roadside units, visible light communication luminaires, free-space optical transceivers, and any future wireless communication apparatus not yet conceived.')
doc.add_paragraph()
para('These movable nodes are capable of physical displacement in any combination of spatial dimensions and/or rotational axes through any actuation mechanism\u2014existing, under development, or to be developed in the future\u2014including but not limited to: motorized rails, robotic wheels, tracks, legs, omnidirectional drives, cable-driven systems, gantry systems, linear actuators, telescopic masts, hydraulic/pneumatic lifts, magnetic levitation, electromagnetic drives, propeller-based flight (drones/UAVs), buoyancy-based systems, peristaltic drives, shape-memory alloys, piezoelectric actuators, MEMS actuators, soft robotics actuators, and any other mechanism that causes physical translation or rotation of the node in space.')
doc.add_paragraph()
para('The system employs an optimization-driven controller\u2014using any optimization algorithm, decision process, heuristic, machine learning model, or control policy (existing, under development, or to be developed)\u2014that senses or estimates the current state of the environment and determines the optimal combination of:')
lp('Physical position (in any number of spatial dimensions and/or rotational axes) of each node.', bold_prefix='\u2022')
lp('Electronic transmission parameters (including but not limited to: beamwidth, beam direction/azimuth/elevation, transmit power, frequency/channel, modulation scheme, coding rate, MIMO configuration, polarization, duty cycle, and any other configurable parameter) of each node.', bold_prefix='\u2022')
lp('The decision of whether to move, how far to move, and when to move each node versus merely adapting its electronic parameters\u2014structured as a hierarchical, multi-tier, or multi-stage decision process that escalates from least disruptive (electronic-only) to most disruptive (full spatial relocation) actions, with each stage invoked only when prior stages are insufficient to meet the coverage objective.', bold_prefix='\u2022')
doc.add_paragraph()

para('Briefly summarize the \u201cunique\u201d or \u201cnovel\u201d or \u201cinventive\u201d aspects of your invention? [Min 50 words]', bold=True)
doc.add_paragraph()

para('The fundamental and primary innovation of this invention is the concept that wireless network infrastructure\u2014access points, base stations, routers, relay nodes, repeaters, and any other wireless communication apparatus\u2014should be physically movable and should autonomously relocate themselves in real time in response to dynamic environmental changes, specifically human blockages, crowd movement, user mobility patterns, and any other time-varying obstruction or demand shift. Unlike all existing wireless deployments where infrastructure is permanently fixed and coverage gaps caused by human movement or environmental changes are accepted as unavoidable limitations, this invention treats the physical position of every wireless node as a continuously optimizable variable that adapts to the environment, just as electronic parameters (beamwidth, power, frequency) are adapted today.')
doc.add_paragraph()
para('The core novelty lies in the following aspects:', bold=True)
doc.add_paragraph()

para('A. Physical Mobility as a First-Class Network Optimization Dimension \u2014 The Primary Innovation:', bold=True)
para('The central inventive concept is that wireless transmitting nodes physically move\u2014translate along X, Y, and/or Z axes, and/or rotate\u2014to restore, maintain, or improve wireless coverage when human bodies, crowd formations, moving equipment, or any other dynamic obstruction degrades signal quality. This is fundamentally different from all prior art, which treats node positions as fixed after initial deployment. The invention recognizes that:')
lp('Human blockage is the dominant and most unpredictable cause of coverage degradation in high-frequency wireless networks. A single human body attenuates signals by 15\u201325 dB at THz/mmWave frequencies, and crowd formations create spatially correlated blockage zones that no fixed-position electronic adaptation can overcome.', bold_prefix='\u2022')
lp('The spatial geometry between a transmitting node and its users is the single most impactful factor in signal quality. Even small physical displacements (centimeters to meters) of a node can dramatically change line-of-sight availability, bypass blockages, reduce path loss, and restore coverage\u2014effects that no amount of electronic beamforming from a fixed position can achieve when the propagation path is physically obstructed.', bold_prefix='\u2022')
lp('By making wireless infrastructure physically mobile, the network gains an entirely new degree of freedom that transcends the fundamental limitations of fixed-position electronic-only optimization. The node does not merely steer its beam around a blockage\u2014it physically moves to a position where the blockage no longer exists in the propagation path.', bold_prefix='\u2022')
lp('This physical mobility applies to all types of wireless communication nodes (access points, base stations, routers, relays, repeaters, small cells, femtocells, picocells, IoT gateways, mesh nodes, RIS, V2X roadside units, optical transceivers, and any future apparatus), including nodes equipped with multiple antennas, phased antenna arrays, and MIMO (Multiple-Input Multiple-Output) capabilities; all frequency bands (sub-1 GHz through optical); all deployment environments (indoor, outdoor, underground, underwater, aerial, space); and any mobility mechanism (wheels, rails, drones, actuators, magnetic levitation, soft robotics, or any mechanism existing or to be developed).', bold_prefix='\u2022')
doc.add_paragraph()

para('B. Autonomous Reactivity to Dynamic Human Blockages and Environmental Changes:', bold=True)
para('The system continuously monitors the deployment environment for changes in human presence, crowd density, user mobility patterns, equipment movement, and any other dynamic obstruction. When these changes degrade coverage below an operator-defined threshold, the system autonomously determines new optimal physical positions for one or more nodes and commands them to relocate\u2014without any manual intervention, network planning, or site survey. The system reacts to:')
lp('Human crowd formation and dispersal (e.g., shift changes in factories, class schedules in universities, event crowds in stadiums, rush hours in transit stations).', bold_prefix='\u2022')
lp('Individual or group human movement that creates transient or sustained line-of-sight blockages.', bold_prefix='\u2022')
lp('Moving equipment, vehicles, robotic platforms, or any other mobile obstruction in the deployment area.', bold_prefix='\u2022')
lp('Environmental changes (weather, door/window opening/closing, furniture rearrangement, temporary structures).', bold_prefix='\u2022')
lp('Changes in user demand distribution (users moving to different areas of a building or outdoor space).', bold_prefix='\u2022')
doc.add_paragraph()

para('C. Hierarchical Decision Framework for Movement vs. Electronic Adaptation:', bold=True)
para('To avoid unnecessary mechanical actuation, the system structures its response as a hierarchical, multi-tier escalation. Physical movement is invoked only when electronic-only adaptation (beamwidth tuning, azimuth steering, power adjustment, frequency switching) has been attempted and found insufficient to restore coverage. This hierarchy ensures that the cheapest, fastest response is tried first, and the more disruptive physical relocation is used only when truly needed. The number of tiers, their ordering, and the degrees of freedom at each tier are all configurable.')
doc.add_paragraph()

para('D. Coordinated Multi-Node Physical Repositioning:', bold=True)
para('When multiple nodes must move, the system optimizes all node positions jointly as a coordinated collective\u2014considering their mutual interference, spatial distribution, and the collective coverage they provide\u2014rather than moving each node independently. This joint optimization uses any suitable algorithm (swarm-based, evolutionary, gradient-based, reinforcement learning, or any future approach) and evaluates a composite objective that balances coverage quality, density-weighted service, movement cost, inter-node separation, and operator-defined priorities.')
doc.add_paragraph()

para('E. Displacement-Minimizing Physical Assignment:', bold=True)
para('After optimization determines new target positions, the system assigns these positions to existing physical nodes using a displacement-minimizing mapping (e.g., Hungarian algorithm or any assignment algorithm) to ensure each node travels the shortest possible distance, minimizing transit time, energy expenditure, mechanical wear, and service interruption during the relocation.')
doc.add_paragraph()

para('Enlist all advantages of the invention.', bold=True)
doc.add_paragraph()
advs = [
    ('Sustained Coverage:', 'Maintains any operator-defined coverage threshold despite dynamic changes in user density, blockages, and environment.'),
    ('Low Mechanical Latency:', 'Hierarchical escalation ensures physical movement only when electronic adaptation is insufficient, reducing actuation delay and wear.'),
    ('Autonomous Coordination:', 'All nodes adapt collectively via optimization, finding globally optimal configurations without manual intervention.'),
    ('Resource Efficiency:', 'Fewer total nodes needed versus static deployments\u2014existing nodes physically reposition to serve shifting demand.'),
    ('Universal Node Applicability:', 'Applies to any wireless communication node type\u2014access points, base stations, routers, relays, repeaters, RIS, IoT gateways, mesh nodes, V2X units, optical transceivers, and any future node type.'),
    ('Universal Frequency Applicability:', 'Operates across sub-1 GHz, sub-6 GHz, mmWave, THz, infrared, optical, visible light, and any future frequency band.'),
    ('Universal Algorithm Applicability:', 'Compatible with any optimization algorithm\u2014PSO, GA, DE, ACO, RL, Bayesian, gradient-based, rule-based, hybrid, or any future algorithm.'),
    ('Universal Environment Applicability:', 'Works in indoor, outdoor, underground, underwater, aerial, space, or any mixed environment.'),
    ('Universal Mobility Mechanism:', 'Supports any physical mobility mechanism\u2014wheels, rails, drones, gantries, actuators, magnetic levitation, soft robotics, or any future mechanism.'),
    ('Operator-Configurable:', 'All objectives, thresholds, tier structures, weights, and constraints are operator-configurable.'),
    ('Minimal Service Interruption:', 'Displacement-minimizing assignment ensures shortest transit per node.'),
    ('Material-Aware and 3D-Aware:', 'Propagation model incorporates material-specific attenuation, height-aware ray tracing, and 3D geometry.'),
    ('Scalable:', 'Scales to any number of nodes and any deployment area size.'),
    ('Future-Proof:', 'Framework is algorithm-agnostic, node-agnostic, frequency-agnostic, and mechanism-agnostic, accommodating technologies not yet invented.'),
]
for t, d in advs:
    lp(d, bold_prefix=t)
doc.add_paragraph()

para('Testing (Very Important)', bold=True)
para('Has the invention been tested experimentally or prototypes? \u2013 If yes, please provide detailed experimental details showing the efficacy of the invention.', bold=True)
doc.add_paragraph()
para('Yes, the invention has been extensively tested via robust simulation models.')
doc.add_paragraph()
lp('A 3D discrete grid of 32 m \u00d7 20 m \u00d7 3 m room layout with varying material attenuation (Concrete: 30 dB, Wood: 20 dB, Glass: 15 dB). Multiple wall types including solid structural walls, glass partitions, wood partitions, and shorter obstacles (1 m height).', bold_prefix='Simulation Environment:')
lp('Evaluated at 300 GHz carrier frequency, 20 dBm transmit power, 30 dBi max antenna gain, \u221294 dBm noise floor. Beamwidth candidates: 30\u00b0, 45\u00b0, 60\u00b0, 75\u00b0, 90\u00b0, 120\u00b0. Azimuth candidates every 45\u00b0. Z-axis heights 2.0\u20133.0 m (0.2 m steps). 7 nodes deployed. SINR threshold: 5 dB. 13 human density clusters per time step (radius 2.0 m each, 20 dB blockage penalty per person). Swarm: 20 particles, 20 iterations per tier. Movement penalty weight: 0.15. Anti-clustering weight: 0.3. Min inter-node distance: 3.0 m. Density weight: 0.6, Cell weight: 0.4.', bold_prefix='Setup:')
lp('The 4-tier hierarchical optimization sustains coverage above the configurable target (tested at 95% density-weighted stringency, generalizable to any 50\u201399.9% threshold) across multiple time steps. Before/after SINR heatmaps show significant improvement. Total movement minimized via tier escalation and Hungarian assignment. Computation time: seconds per cycle. Reachability increase per unit movement demonstrates efficient mechanical utilization.', bold_prefix='Results:')
doc.add_paragraph()

heading2('Please attach illustrations (architecture diagrams / flowcharts / block diagrams / product illustrations etc.)')
doc.add_paragraph()
for item in [
    'Flowchart 1: Hierarchical multi-tier optimization flow with decision points for tier escalation.',
    'Graph 1: Optimization time vs. time step.',
    'Graph 2: Density-weighted coverage vs. time step.',
    'Graph 3: Total node movement (m) vs. time step.',
    'Graph 4: Population density heatmap at each time step.',
    'Graph 5: SINR heatmap before and after optimization.',
    'Graph 6: Reachability increase vs. movement (efficiency plot).',
    'Graph 7: Node placement maps with directional indicators before/after optimization.',
    'FIG. A: System architecture block diagram.',
    'FIG. B: Multiple mechanical mobility platform embodiments.',
    'FIG. C: 3D propagation model illustration.',
    'FIG. D: Composite fitness function computation illustration.',
]:
    para(item); doc.add_paragraph()

lp('Commercial potential', bold_prefix='')
doc.add_paragraph()
para('7.1 Please list possible uses or application areas, or products that may embody some aspects of the technology:', bold=True)
for a in [
    '6G/5G indoor network infrastructure for smart factories, warehouses, and stadia.',
    'AR/VR/XR gaming arenas and immersive venues requiring ultra-high-bandwidth, ultra-low-latency feeds.',
    'Dense enterprise offices with dynamic seating/hot-desking.',
    'Hospitals/healthcare with shift-driven density changes.',
    'Retail/malls with time-varying customer flow.',
    'Educational institutions with schedule-driven density.',
    'Data centers for intra-facility wireless management.',
    'Industry 4.0 plants with robotic platforms.',
    'Self-optimizing IoT base station networks for smart cities.',
    'Military/defense tactical networks.',
    'Emergency response/disaster relief temporary deployments.',
    'Outdoor urban: street canyons, parks, transit stations.',
    'Underground: mines, tunnels, subways.',
    'Maritime/offshore: ships, oil platforms.',
    'Agricultural: precision farming connectivity.',
    'Construction sites with changing layouts.',
    'Airport terminals and aircraft hangars.',
    'Large-scale event venues (concerts, exhibitions, sports).',
]:
    lp(a, bold_prefix='\u2022')
doc.add_paragraph()
para('7.2 List of probable users of the technology:', bold=True)
para('Telecom operators (Verizon, AT&T, Jio, Vodafone, T-Mobile, China Mobile, NTT Docomo, SK Telecom), Equipment vendors (Ericsson, Nokia, Samsung, Huawei, ZTE, Qualcomm, Intel, Cisco, CommScope), Smart building integrators, Industry 4.0 companies, Cloud/edge providers, Defense agencies, Robotics companies, IoT platform providers, and any entity deploying wireless infrastructure.')
doc.add_paragraph()

para('Prior self publication:', bold=True)
para('Please declare any self publications made in any academic paper / journal / scientific conference etc.', bold=True)
pt = doc.add_table(rows=2, cols=4); pt.style = 'Table Grid'
for i, h in enumerate(['Disclosure Type', 'Date', 'Citation', 'URL']):
    c = pt.rows[0].cells[i]; c.text = h
    for p in c.paragraphs:
        for r in p.runs: r.bold = True; r.font.size = Pt(10)
for i in range(4): pt.rows[1].cells[i].text = '[To be filled]'
doc.add_paragraph()

para('Application Country(ies):', bold=True)
para('India, United States, European Union, Japan, South Korea, China.')
doc.add_paragraph()

doc.add_page_break()
heading1('ADDITIONAL SHEET: EXHAUSTIVE PATENT CLAIMS')
doc.add_paragraph()
heading2('Independent Claims')
doc.add_paragraph()

para('Claim 1. A system for dynamically optimizing wireless network coverage, comprising:', bold=True)
for item in [
    '(a) one or more movable wireless communication nodes, wherein each node is any device or apparatus\u2014existing, under development, or to be developed in the future\u2014capable of at least one of: transmitting, receiving, relaying, repeating, amplifying, reflecting, refracting, scattering, or otherwise processing wireless electromagnetic signals at any frequency or combination of frequencies, including but not limited to sub-1 GHz, sub-6 GHz, millimeter-wave (24\u2013100 GHz), terahertz (100 GHz\u201310 THz), infrared, optical, and visible light frequencies; and including but not limited to: access points, base stations (macro, micro, pico, femto, nano), routers, relay nodes, repeaters, range extenders, distributed antenna system heads, radio units, remote radio heads, reconfigurable intelligent surfaces with active or passive elements, intelligent reflecting surfaces, backhaul nodes, fronthaul nodes, mesh network nodes, ad-hoc network nodes, mobile hotspots, Wi-Fi extenders, IoT gateways, LoRa gateways, satellite ground terminals, vehicle-to-everything (V2X) roadside units, visible light communication luminaires, free-space optical transceivers, and any future wireless communication apparatus;',
    '(b) each said movable node comprising: (i) a physical mobility mechanism\u2014of any type existing, under development, or to be developed\u2014configured to change the spatial position and/or orientation of the node in at least one degree of freedom, including but not limited to: translational movement along any combination of a first horizontal axis (X), a second horizontal axis (Y), a vertical axis (Z), and/or any arbitrary axis; and/or rotational movement about any combination of roll, pitch, yaw, or any arbitrary axis; said mechanism including but not limited to: motorized wheels, omnidirectional drives, tracked drives, legged locomotion, motorized rails or tracks, cable-driven systems, gantry systems, linear actuators, telescopic masts, hydraulic lifts, pneumatic lifts, magnetic levitation, electromagnetic linear motors, propeller-based flight, rotor-based flight, buoyancy-based vertical adjustment, shape-memory alloy actuators, piezoelectric actuators, MEMS actuators, soft robotics actuators, peristaltic drives, any robotic platform, and any other mechanism that causes physical displacement or rotation of the node; (ii) one or more adjustable transmission or reception parameters, including but not limited to: beamwidth, beam direction (azimuth angle, elevation angle, or any angular coordinate), transmit power level, receive sensitivity, frequency or channel selection, bandwidth, modulation scheme, coding rate, MIMO antenna configuration, polarization, duty cycle, beam pattern shape, sidelobe control, null steering, and any other electronically, mechanically, or software-configurable parameter that affects the radiation pattern, signal quality, coverage, capacity, or interference characteristics of the node;',
    '(c) one or more sensing, estimation, or inference mechanisms\u2014of any type existing, under development, or to be developed\u2014configured to determine or approximate the state of the deployment environment, including but not limited to: the spatial distribution of dynamic blockages (humans, vehicles, equipment, animals, weather phenomena, or any other time-varying obstruction); user density or traffic demand distribution; signal quality measurements (received signal strength, SINR, SNR, RSRP, RSRQ, CQI, throughput, latency, jitter, packet loss, bit error rate, or any other quality metric); static infrastructure characteristics (wall locations, material types, attenuation values, terrain, building geometry); and any other environmental variable relevant to wireless coverage; said sensing mechanisms including but not limited to: cameras (RGB, depth, stereo, 360-degree), LiDAR, radar (mmWave, UWB), ultrasonic sensors, thermal/infrared sensors, occupancy sensors, motion detectors, pressure sensors, environmental sensors, Wi-Fi/Bluetooth/RF fingerprinting, signal strength measurements, channel state information, crowdsourced data, user equipment reports, GPS/GNSS positioning, inertial measurement units, digital twin models, predictive/forecasting models (statistical, machine learning, deep learning), facility management system feeds, building information modeling (BIM) data, and any other data source or sensor modality;',
    '(d) a controller or decision-making entity\u2014centralized, distributed, hierarchical, peer-to-peer, cloud-based, edge-based, or any combination thereof\u2014configured to: (i) determine, using any computational method (including but not limited to: algorithmic optimization, heuristic rules, machine learning inference, reinforcement learning policy, expert systems, fuzzy logic, model predictive control, digital twin simulation, or any other existing or future computational approach), a target configuration comprising the spatial position, orientation, and/or transmission parameters of each of the one or more movable nodes; (ii) evaluate the target configuration against an operator-defined or autonomously-determined coverage objective, quality-of-service objective, or performance objective (expressed as any metric, threshold, or constraint, including but not limited to: percentage of area covered, percentage of user density served, minimum SINR/SNR/RSRP at any point, maximum latency, minimum throughput, maximum bit error rate, fairness index, energy efficiency, or any other performance criterion set to any numerical value from 0% to 100% or any other range); (iii) structure the optimization, search, or decision process as a hierarchical, multi-tier, multi-stage, or multi-phase procedure\u2014with any number of tiers or stages (including but not limited to 2, 3, 4, 5, or more)\u2014wherein earlier stages modify fewer or less disruptive degrees of freedom (e.g., electronic parameters only) and later stages unlock additional, more disruptive degrees of freedom (e.g., vertical repositioning, horizontal repositioning, full 3D repositioning), with each subsequent stage invoked only when prior stages are insufficient to achieve the coverage objective; and the number of stages, ordering of stages, and degrees of freedom at each stage are configurable; (iv) dispatch, communicate, or otherwise cause each movable node to execute its assigned repositioning, reorientation, and/or parameter adjustment.',
]:
    lp(item, bold_prefix='\u2022')

doc.add_paragraph()
para('Claim 2. A method for dynamically optimizing wireless network coverage, comprising:', bold=True)
for item in [
    '(a) providing one or more movable wireless communication nodes as described in Claim 1(a)\u2013(b);',
    '(b) sensing, estimating, or inferring the state of the deployment environment, including the spatial distribution of dynamic blockages and/or user density and/or traffic demand using any sensing or estimation mechanism as described in Claim 1(c);',
    '(c) computing or estimating a signal quality metric or coverage metric at a plurality of spatial locations within the deployment area, based on: the current positions, orientations, and transmission parameters of the one or more movable nodes; attenuation characteristics of static infrastructure; and the sensed or estimated distribution of dynamic blockages;',
    '(d) determining whether the computed coverage metric satisfies a configurable coverage objective;',
    '(e) when the coverage objective is not satisfied, executing a hierarchical, multi-tier, or multi-stage optimization or decision procedure that escalates from electronic-only parameter adaptation through progressively more disruptive physical repositioning stages, with each subsequent stage invoked only when prior stages are insufficient, and where any optimization algorithm, heuristic, machine learning model, or decision process may be used at any stage;',
    '(f) assigning the optimized positions and parameters to the physical movable nodes, optionally using a displacement-minimizing or cost-minimizing assignment algorithm; and',
    '(g) causing each movable node to execute its assigned repositioning, reorientation, and/or parameter adjustment.',
]:
    lp(item, bold_prefix='\u2022')

doc.add_paragraph()
para('Claim 3. A non-transitory computer-readable storage medium storing instructions that, when executed by one or more processors, cause the one or more processors to perform the method of Claim 2.', bold=True)
doc.add_paragraph()
para('Claim 4. A movable wireless communication node apparatus as described in Claim 1(a)\u2013(b), further comprising a communication interface for receiving repositioning commands and a processor for executing a portion of the optimization or relaying sensor data to a controller.', bold=True)
doc.add_paragraph()
para('Claim 5. A wireless communication network comprising a plurality of systems as described in Claim 1, optionally coordinated by a central or distributed management entity across multiple zones, rooms, floors, buildings, or geographic areas.', bold=True)
doc.add_paragraph()

heading2('Dependent Claims')
doc.add_paragraph()

deps = [
    'Claim 6. The system of Claim 1, wherein the one or more carrier frequencies include terahertz frequencies (100 GHz to 10 THz).',
    'Claim 7. The system of Claim 1, wherein the one or more carrier frequencies include millimeter-wave frequencies (24 GHz to 100 GHz).',
    'Claim 8. The system of Claim 1, wherein the one or more carrier frequencies include sub-6 GHz frequencies.',
    'Claim 9. The system of Claim 1, wherein the one or more carrier frequencies include optical or visible light frequencies.',
    'Claim 10. The system of Claim 1, wherein the mobility mechanism provides translational movement in all three spatial dimensions (X, Y, and Z) simultaneously.',
    'Claim 11. The system of Claim 1, wherein the mobility mechanism provides rotational movement about at least one axis in addition to translational movement.',
    'Claim 12. The system of Claim 1, wherein the mobility mechanism comprises at least one of: ceiling-mounted rail, wheeled robot, tracked robot, legged robot, gantry, linear actuator, telescopic mast, hydraulic lift, pneumatic lift, magnetic levitation, drone/UAV, cable-driven platform, shape-memory alloy actuator, piezoelectric actuator, MEMS actuator, soft robotics actuator, or any other mechanism.',
    'Claim 13. The system of Claim 1, wherein the adjustable parameters include at least two of: beamwidth, beam direction, transmit power, frequency selection, modulation, MIMO configuration, polarization, and duty cycle.',
    'Claim 14. The system of Claim 1, wherein beamwidth is adjustable across a range from 1 degree to 360 degrees.',
    'Claim 15. The system of Claim 1, wherein azimuth orientation is adjustable across 0 to 360 degrees.',
    'Claim 16. The system of Claim 1, wherein the sensing mechanism comprises at least one of: camera, LiDAR, radar, ultrasonic sensor, thermal sensor, occupancy sensor, RF fingerprinting, channel state information, crowdsourced data, user equipment reports, digital twin model, predictive AI model, BIM data, or any combination thereof.',
    'Claim 17. The system of Claim 1, wherein the signal quality metric comprises at least one of: SINR, SNR, RSRP, RSRQ, CQI, throughput, latency, jitter, packet loss, bit error rate, packet error rate, spectral efficiency, energy efficiency, or any other measurable quality indicator.',
    'Claim 18. The system of Claim 1, wherein the coverage objective is expressed as a percentage of total user density served, the percentage being any value from 0.1% to 100%.',
    'Claim 19. The system of Claim 1, wherein the coverage objective is expressed as a percentage of total area covered, the percentage being any value from 0.1% to 100%.',
    'Claim 20. The system of Claim 1, wherein the coverage objective is dynamically adjustable based on any combination of: network load, time of day, day of week, event schedules, SLA requirements, energy budget, or any other operational condition.',
    'Claim 21. The system of Claim 1, wherein the controller employs a population-based or swarm-intelligence-based optimization algorithm, including at least one of: Particle Swarm Optimization (and any variant thereof), Genetic Algorithm (and any variant), Differential Evolution, Ant Colony Optimization, Artificial Bee Colony, Grey Wolf Optimizer, Whale Optimization Algorithm, Simulated Annealing, Tabu Search, Cuckoo Search, Firefly Algorithm, Bat Algorithm, Moth-Flame Optimization, Harris Hawks Optimization, Salp Swarm Algorithm, Marine Predators Algorithm, Harmony Search, Gravitational Search, Sine Cosine Algorithm, Aquila Optimizer, or any hybrid, ensemble, or variant thereof.',
    'Claim 22. The system of Claim 1, wherein the controller employs a machine learning or reinforcement learning-based optimization, including at least one of: Deep Q-Network, Proximal Policy Optimization, Actor-Critic, A3C, SAC, TD3, DDPG, model-based RL, Monte Carlo Tree Search, transformer-based decision models, graph neural network-based optimization, multi-agent RL, or any variant.',
    'Claim 23. The system of Claim 1, wherein the controller employs a gradient-based, Bayesian, surrogate-model-based, rule-based, expert system, fuzzy logic, model predictive control, or any deterministic optimization approach.',
    'Claim 24. The system of Claim 1, wherein the controller evaluates candidate configurations using a composite objective function comprising at least two of: (a) coverage ratio; (b) density-weighted coverage; (c) movement/displacement penalty; (d) inter-node separation/anti-clustering penalty; (e) target achievement bonus; (f) energy consumption term; (g) fairness term; (h) latency term; (i) interference minimization term; (j) load balancing term; with configurable weights for each.',
    'Claim 25. The system of Claim 24, wherein every component and weight of the composite objective function is independently configurable by an operator or autonomously adjustable by the system.',
    'Claim 26. The system of Claim 1, wherein the hierarchical procedure comprises at least two tiers, wherein the first tier modifies only electronic transmission parameters and the second tier additionally modifies at least one physical degree of freedom.',
    'Claim 27. The system of Claim 1, wherein the hierarchical procedure comprises at least three tiers: a first tier modifying electronic parameters only; a second tier additionally modifying vertical position (Z-axis); and a third tier additionally modifying horizontal positions (X and/or Y axes).',
    'Claim 28. The system of Claim 1, wherein the hierarchical procedure comprises four or more tiers, including tiers for: beamwidth adaptation, azimuth adaptation, vertical repositioning, horizontal repositioning, power adaptation, frequency adaptation, and/or any other parameter, in any configurable order.',
    'Claim 29. The system of Claim 1, wherein the number, order, and scope of tiers in the hierarchical procedure are configurable by the operator or autonomously adjustable by the system.',
    'Claim 30. The system of Claim 1, wherein at each tier, dimensions or parameters not being optimized are locked (their update velocities, gradients, or perturbations set to zero), and unlocked in subsequent tiers.',
    'Claim 31. The system of Claim 1, wherein optimized positions are assigned to physical nodes using a displacement-minimizing or cost-minimizing assignment algorithm, including at least one of: Hungarian algorithm, auction algorithm, linear sum assignment, shortest augmenting path, or any other assignment/matching algorithm.',
    'Claim 32. The system of Claim 31, wherein the cost matrix is based on any distance metric including: Euclidean distance (2D or 3D), Manhattan distance, weighted distance, travel time, energy cost, or any custom cost function.',
    'Claim 33. The system of Claim 1, wherein each candidate solution is encoded as a vector of dimensionality k x d, where k is the number of nodes and d is the number of adjustable parameters per node, and the total dimensionality scales linearly with the number of nodes.',
    'Claim 34. The system of Claim 21, wherein a configurable fraction of initial candidate solutions are seeded near current node positions and/or configurations, and a remaining fraction are initialized randomly or quasi-randomly across the search space.',
    'Claim 35. The system of Claim 1, wherein the signal quality computation uses a propagation model accounting for at least two of: free-space path loss, atmospheric/molecular absorption, material-specific wall/obstacle attenuation, dynamic blockage attenuation, multipath fading, diffraction, reflection, scattering, and/or any other propagation phenomenon.',
    'Claim 36. The system of Claim 35, wherein the propagation model performs three-dimensional ray-segment or ray-tracing intersection tests considering the height, position, material, and geometry of each obstacle.',
    'Claim 37. The system of Claim 35, wherein dynamic blockage attenuation is applied only at path segments where the ray height is below a blockage height threshold.',
    'Claim 38. The system of Claim 35, wherein material attenuation values are configurable per material type and per frequency band.',
    'Claim 39. The system of Claim 1, wherein the dynamic blockage or user density distribution is represented as a spatial matrix or field D(x, y, t), D(x, y, z, t), or any other spatial-temporal representation, updated at configurable intervals or event-triggered.',
    'Claim 40. The system of Claim 39, wherein the density representation is generated from: real-time sensor fusion, time-series forecasting, Gaussian process regression, deep learning prediction, scheduled profiles, historical analytics, digital twin simulation, or any combination thereof.',
    'Claim 41. The system of Claim 39, wherein each density source generates a spatial kernel (linear, Gaussian, polynomial, or any radial/non-radial basis function) with configurable extent.',
    'Claim 42. The system of Claim 1, wherein the deployment area is any of: indoor, outdoor, underground, underwater, aerial, orbital/space, or any mixed environment.',
    'Claim 43. The system of Claim 1, wherein the deployment area spans multiple rooms, floors, buildings, or geographic zones, with optimization coordinated across zones.',
    'Claim 44. The system of Claim 1, wherein the movable nodes comprise at least one of: access points, base stations (macro/micro/pico/femto/nano), routers, relay nodes, repeaters, range extenders, DAS heads, radio units, remote radio heads, RIS with active elements, backhaul nodes, fronthaul nodes, mesh nodes, ad-hoc nodes, mobile hotspots, Wi-Fi extenders, IoT gateways, LoRa gateways, satellite ground terminals, V2X roadside units, visible light luminaires, free-space optical transceivers, or any future wireless apparatus.',
    'Claim 45. The system of Claim 1, wherein the controller periodically re-evaluates and re-triggers optimization at configurable intervals (ranging from sub-second to hours) or in response to detected environmental changes, user mobility events, or operator commands.',
    'Claim 46. The system of Claim 1, wherein the controller computes and reports at least one of: total displacement per cycle, optimization time, coverage improvement, density-weighted reachability increase, per-node displacement, energy consumed, service interruption duration, or any other performance metric.',
    'Claim 47. The system of Claim 1, wherein all movable nodes are optimized jointly as a coordinated collective using a single optimization objective, rather than independently.',
    'Claim 48. The system of Claim 1, wherein nodes at zone boundaries participate in optimization of adjacent zones for seamless inter-zone coverage continuity.',
    'Claim 49. The system of Claim 1, wherein the controller is implemented as any of: a centralized server, distributed across the nodes, cloud-based, edge-based, hierarchical (with local and global controllers), peer-to-peer consensus, or any combination thereof.',
    'Claim 50. The method of Claim 2, wherein the optimization at each stage uses a different algorithm or the same algorithm with different hyperparameters.',
    'Claim 51. The method of Claim 2, wherein the method is executed continuously in a closed-loop fashion, with the environment re-sensed after each optimization cycle and the process repeated indefinitely.',
    'Claim 52. The method of Claim 2, wherein the method is executed in response to discrete events including: user density threshold exceedance, blockage detection, time schedule trigger, operator command, or any external event.',
    'Claim 53. The method of Claim 2, wherein the coverage objective and optimization parameters are adjusted in real time based on feedback from prior optimization cycles (adaptive/self-tuning optimization).',
    'Claim 54. The apparatus of Claim 4, wherein the node is self-powered (battery, solar, energy harvesting, fuel cell, wireless power transfer, or any power source) and autonomously navigates using onboard SLAM, visual odometry, or any localization method.',
    'Claim 55. The apparatus of Claim 4, wherein the node comprises multiple antennas, antenna arrays, or reconfigurable antenna elements capable of simultaneous multi-beam or MIMO operation.',
    'Claim 56. The network of Claim 5, wherein the network integrates with existing cellular infrastructure (4G/5G/6G core network) and the movable nodes register as standard network elements.',
    'Claim 57. The network of Claim 5, wherein the network operates as a standalone, self-organizing network independent of cellular infrastructure.',
    'Claim 58. The system of Claim 1, wherein any combination of the features described in Claims 6\u201357 are combined in any permutation.',
    'Claim 59. The method of Claim 2, wherein any combination of the features described in Claims 6\u201357 are incorporated.',
    'Claim 60. Any system, method, apparatus, or network that performs dynamic physical repositioning and/or reorientation of one or more wireless communication nodes in response to environmental sensing or estimation, using any form of optimization or decision process that includes at least one stage of electronic parameter adaptation and at least one stage of physical repositioning, substantially as described herein.',
]
for dc in deps: para(dc)

doc.add_paragraph()
heading2('Additional Independent Claim')
doc.add_paragraph()

para('Claim 61. A method for mitigating dynamic human blockages and environmental obstructions in a wireless communication network, the method comprising:', bold=True)
for item in [
    '(a) providing a wireless communication node comprising a multi-antenna array capable of Multiple-Input Multiple-Output (MIMO) operation and electronic beamforming, wherein the node is mechanically coupled to a mobility mechanism configured to provide physical translation in at least one spatial dimension;',
    '(b) detecting a degradation in a wireless signal metric below a target threshold, wherein said degradation is caused by a physical blockage obstructing a propagation path;',
    '(c) attempting to restore the wireless signal metric above the target threshold entirely via electronic means by adjusting one or more electronic transmission parameters, said adjusting comprising steering a beam emitted by the multi-antenna array or altering a MIMO configuration;',
    '(d) determining that said adjusting of electronic transmission parameters is insufficient to restore the wireless signal metric due to the physical blockage; and',
    '(e) in response to said determining, autonomously executing physical translation of the wireless communication node via the mobility mechanism to a computed new spatial position wherein the physical blockage no longer obstructs the propagation path, thereby restoring the wireless signal metric.'
]:
    lp(item, bold_prefix='\u2022')

doc.add_paragraph()
para('Attach additional sheets wherever needed.', bold=True)
doc.save('/home/arpit/Desktop/kingAB/coding/APSO_SIM_Thz/Patent_Draft_Complete.docx')
print('DONE')
